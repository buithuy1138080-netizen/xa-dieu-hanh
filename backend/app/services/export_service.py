"""Export Service — generates DOCX and XLSX report files."""
from __future__ import annotations

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

from app.core.config import settings as _settings
_EXPORT_DIR = Path(_settings.UPLOAD_DIR) / "reports"
_EXPORT_DIR.mkdir(parents=True, exist_ok=True)

# ── DOCX ───────────────────────────────────────────────────────────────────────

def export_docx(report_id: int, title: str, period_label: str,
                summary_data: dict, ai_summary: dict) -> str:
    """Generate a DOCX file, save to disk, return the file path string."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────────
    for section in doc.sections:
        from docx.shared import Cm
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    def _center(text: str, bold: bool = False, size: int = 12) -> None:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    def _heading(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def _para(text: str) -> None:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Pt(24)

    # ── Header ──────────────────────────────────────────────────────────────
    _center("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=13)
    _center("Độc lập – Tự do – Hạnh phúc", bold=False, size=12)
    _center("─────────────────────", bold=False, size=10)
    doc.add_paragraph()

    # ── Title ────────────────────────────────────────────────────────────────
    _center(title.upper(), bold=True, size=14)
    _center(period_label, bold=False, size=12)
    doc.add_paragraph()

    # ── I. Executive Summary ─────────────────────────────────────────────────
    _heading("I. TỔNG QUAN CHUNG")
    _para(ai_summary.get("tong_quat", ""))

    _heading("II. ĐÁNH GIÁ TIẾN ĐỘ THỰC HIỆN")
    _para(ai_summary.get("danh_gia_tien_do", ""))

    # ── II. Task table ───────────────────────────────────────────────────────
    tasks = summary_data.get("tasks", {})
    if tasks.get("total"):
        _heading("III. KẾT QUẢ THỰC HIỆN NHIỆM VỤ", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Chỉ tiêu"
        hdr[1].text = "Số lượng"
        hdr[2].text = "Tỷ lệ"
        rows_data = [
            ("Tổng nhiệm vụ",    str(tasks["total"]),     ""),
            ("Hoàn thành",       str(tasks["completed"]), f"{tasks['completion_rate']}%"),
            ("Đang thực hiện",   str(tasks["in_progress"]), ""),
            ("Chờ xử lý",        str(tasks["pending"]), ""),
            ("Quá hạn",          str(tasks["overdue"]), ""),
        ]
        for d in rows_data:
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text = d

    # ── III. KPI summary ─────────────────────────────────────────────────────
    kpis = summary_data.get("kpis", {})
    if kpis.get("total"):
        doc.add_paragraph()
        _heading("IV. KẾT QUẢ CHỈ TIÊU KPI CHIẾN LƯỢC", level=2)
        doc.add_paragraph(
            f"Tổng số KPI: {kpis['total']} | Bình quân hoàn thành: {kpis['avg_pct']}%"
        )

    # ── Overdue tasks ─────────────────────────────────────────────────────────
    overdue = summary_data.get("overdue_tasks", [])
    if overdue:
        doc.add_paragraph()
        _heading("V. NHIỆM VỤ QUÁ HẠN (CẦN XỬ LÝ)")
        table2 = doc.add_table(rows=1, cols=4)
        table2.style = "Table Grid"
        h = table2.rows[0].cells
        h[0].text, h[1].text, h[2].text, h[3].text = "Tên nhiệm vụ", "Đơn vị", "Hạn", "Ngày trễ"
        for t in overdue[:10]:
            r = table2.add_row().cells
            r[0].text = t.get("title", "")[:60]
            r[1].text = t.get("dept", "—")
            r[2].text = (t.get("due_date") or "—")[:10]
            r[3].text = str(t.get("days_late", 0))

    # ── VI. Conclusions ───────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading("VI. TỒN TẠI, HẠN CHẾ VÀ NGUYÊN NHÂN")
    for line in ai_summary.get("ton_tai_han_che", "").splitlines():
        if line.strip():
            doc.add_paragraph(line.strip(), style="List Bullet")
    doc.add_paragraph("Nguyên nhân:").runs[0].bold = True
    for line in ai_summary.get("nguyen_nhan", "").splitlines():
        if line.strip():
            doc.add_paragraph(line.strip(), style="List Bullet")

    _heading("VII. KIẾN NGHỊ VÀ NHIỆM VỤ TRỌNG TÂM")
    for line in ai_summary.get("kien_nghi", "").splitlines():
        if line.strip():
            doc.add_paragraph(line.strip(), style="List Bullet")
    doc.add_paragraph("Nhiệm vụ trọng tâm tiếp theo:").runs[0].bold = True
    for line in ai_summary.get("nhiem_vu_trong_tam", "").splitlines():
        if line.strip():
            doc.add_paragraph(line.strip(), style="List Bullet")

    # ── Signature block ───────────────────────────────────────────────────────
    doc.add_paragraph()
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.rows[0].cells[0].text = (
        f"Nơi nhận:\n- Như trên;\n- Lưu VT.\n\n"
        f"Ngày {datetime.now().strftime('%d tháng %m năm %Y')}"
    )
    sig_tbl.rows[0].cells[1].text = (
        "TM. ỦY BAN NHÂN DÂN\nCHỦ TỊCH\n\n\n\n(Ký, đóng dấu)"
    )
    for cell in sig_tbl.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Save ─────────────────────────────────────────────────────────────────
    fname = f"report_{report_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    fpath = _EXPORT_DIR / fname
    doc.save(str(fpath))
    return str(fpath)


# ── XLSX ───────────────────────────────────────────────────────────────────────

def export_xlsx(report_id: int, title: str, period_label: str,
                summary_data: dict, ai_summary: dict) -> str:
    """Generate an XLSX file, return the file path string."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise RuntimeError("openpyxl not installed. Run: pip install openpyxl")

    wb = Workbook()

    # ── Sheet 1: Summary ─────────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Tổng quan"
    _hdr_style = Font(bold=True, size=12, color="FFFFFF")
    _hdr_fill  = PatternFill("solid", fgColor="4F46E5")
    _bold      = Font(bold=True)

    def _hrow(ws_: Any, row: int, values: list, fill_color: str = "4F46E5") -> None:
        for col, val in enumerate(values, 1):
            cell = ws_.cell(row=row, column=col, value=val)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=fill_color)
            cell.alignment = Alignment(horizontal="center")

    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = period_label
    ws["A2"].font = Font(italic=True, size=11)
    ws.merge_cells("A1:E1")
    ws.merge_cells("A2:E2")

    # Task stats table
    row = 4
    tasks = summary_data.get("tasks", {})
    _hrow(ws, row, ["CHỈ TIÊU NHIỆM VỤ", "SỐ LƯỢNG", "TỶ LỆ"])
    row += 1
    for label_text, val_key, pct_key in [
        ("Tổng số nhiệm vụ",    "total",      None),
        ("Hoàn thành",          "completed",  "completion_rate"),
        ("Đang thực hiện",      "in_progress", None),
        ("Chờ xử lý",           "pending",    None),
        ("Quá hạn",             "overdue",    None),
    ]:
        ws.cell(row=row, column=1, value=label_text)
        ws.cell(row=row, column=2, value=tasks.get(val_key, 0))
        if pct_key:
            ws.cell(row=row, column=3, value=f"{tasks.get(pct_key, 0)}%")
        row += 1

    # KPI stats
    row += 1
    kpis = summary_data.get("kpis", {})
    if kpis.get("total"):
        _hrow(ws, row, ["CHỈ TIÊU KPI", "SỐ LƯỢNG", "BÌNH QUÂN %"])
        row += 1
        by_st = kpis.get("by_status", {})
        for label_text, val in [
            ("Tổng KPI",         kpis.get("total", 0)),
            ("Bình quân hoàn thành", f"{kpis.get('avg_pct', 0)}%"),
            ("Đạt mục tiêu",     by_st.get("dat_muc_tieu", 0)),
            ("Đúng tiến độ",     by_st.get("dung_tien_do", 0)),
            ("Chậm tiến độ",     by_st.get("cham_tien_do", 0)),
            ("Quá hạn",          by_st.get("qua_han", 0)),
        ]:
            ws.cell(row=row, column=1, value=label_text)
            ws.cell(row=row, column=2, value=val)
            row += 1

    # Column widths
    ws.column_dimensions["A"].width = 35
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 16

    # ── Sheet 2: Department breakdown ────────────────────────────────────────
    ws2 = wb.create_sheet("Phân theo đơn vị")
    _hrow(ws2, 1, ["ĐƠN VỊ", "TỔNG NV", "HOÀN THÀNH", "TỶ LỆ"])
    for i, dept in enumerate(summary_data.get("dept_breakdown", []), 2):
        ws2.cell(i, 1, dept.get("name", ""))
        ws2.cell(i, 2, dept.get("total", 0))
        ws2.cell(i, 3, dept.get("completed", 0))
        ws2.cell(i, 4, f"{dept.get('rate', 0)}%")
    ws2.column_dimensions["A"].width = 25

    # ── Sheet 3: Overdue tasks ────────────────────────────────────────────────
    ws3 = wb.create_sheet("NV Quá hạn")
    _hrow(ws3, 1, ["STT", "TÊN NHIỆM VỤ", "ĐƠN VỊ", "HẠN CHÓT", "SỐ NGÀY TRỄ", "ƯU TIÊN"])
    for i, t in enumerate(summary_data.get("overdue_tasks", [])[:50], 2):
        ws3.cell(i, 1, i - 1)
        ws3.cell(i, 2, t.get("title", "")[:80])
        ws3.cell(i, 3, t.get("dept", ""))
        ws3.cell(i, 4, (t.get("due_date") or "")[:10])
        ws3.cell(i, 5, t.get("days_late", 0))
        ws3.cell(i, 6, t.get("priority", ""))
    ws3.column_dimensions["B"].width = 50

    # ── Sheet 4: AI Summary ───────────────────────────────────────────────────
    ws4 = wb.create_sheet("Nhận xét AI")
    sections = [
        ("TỔNG QUAN CHUNG",          ai_summary.get("tong_quat", "")),
        ("ĐÁNH GIÁ TIẾN ĐỘ",        ai_summary.get("danh_gia_tien_do", "")),
        ("TỒN TẠI HẠN CHẾ",         ai_summary.get("ton_tai_han_che", "")),
        ("NGUYÊN NHÂN",              ai_summary.get("nguyen_nhan", "")),
        ("KIẾN NGHỊ",                ai_summary.get("kien_nghi", "")),
        ("NHIỆM VỤ TRỌNG TÂM",       ai_summary.get("nhiem_vu_trong_tam", "")),
    ]
    r = 1
    for sec_title, content in sections:
        ws4.cell(r, 1, sec_title).font = Font(bold=True, size=11)
        r += 1
        for line in content.splitlines():
            if line.strip():
                c = ws4.cell(r, 1, line.strip())
                c.alignment = Alignment(wrap_text=True)
                r += 1
        r += 1
    ws4.column_dimensions["A"].width = 100

    # ── Save ─────────────────────────────────────────────────────────────────
    fname = f"report_{report_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    fpath = _EXPORT_DIR / fname
    wb.save(str(fpath))
    return str(fpath)
