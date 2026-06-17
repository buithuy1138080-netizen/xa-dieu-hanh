"""
template_engine.py

Renders .xlsx and .docx templates by substituting {{variable}} placeholders.

Syntax:
  Scalar  : {{var_name}}          — replaced with scalar value from variables dict
  Loop    : {{#list_var}}         — start marker row/paragraph
              {{item.field}}      — per-item fields inside loop
            {{/list_var}}         — end marker row/paragraph

PDF: exported via DOCX → LibreOffice conversion (requires libreoffice on PATH).
If LibreOffice unavailable, DOCX is returned instead.
"""
from __future__ import annotations

import copy
import logging
import re
import subprocess
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

VAR_RE = re.compile(r"\{\{(\w+(?:\.\w+)?)\}\}")
LOOP_START_RE = re.compile(r"^\s*\{\{#(\w+)\}\}\s*$")
LOOP_END_RE   = re.compile(r"^\s*\{\{/(\w+)\}\}\s*$")


# ── Public API ─────────────────────────────────────────────────────────────────

def render_xlsx(template_path: str, variables: dict[str, Any], output_path: str) -> None:
    import openpyxl
    wb = openpyxl.load_workbook(template_path)
    for ws in wb.worksheets:
        _render_xlsx_sheet(ws, variables)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)


def render_docx(template_path: str, variables: dict[str, Any], output_path: str) -> None:
    from docx import Document
    doc = Document(template_path)
    _render_docx(doc, variables)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def render_pdf(template_path: str, variables: dict[str, Any], output_path: str) -> str:
    """
    Render template to PDF. Strategy:
    1. Render to DOCX first.
    2. Convert DOCX → PDF via LibreOffice (headless).
    3. If LibreOffice not found, return DOCX path instead.
    Returns the actual output path (may differ if LibreOffice unavailable).
    """
    docx_path = output_path.replace(".pdf", "_tmp.docx")
    render_docx(template_path, variables, docx_path)

    pdf_path = _convert_docx_to_pdf(docx_path, output_path)
    return pdf_path


def parse_variables(template_path: str) -> tuple[list[str], list[str]]:
    """
    Scan a template file and return (scalar_vars, list_vars) found in it.
    Works for both .xlsx and .docx.
    """
    ext = Path(template_path).suffix.lower().lstrip(".")
    if ext == "xlsx":
        return _parse_variables_xlsx(template_path)
    if ext == "docx":
        return _parse_variables_docx(template_path)
    return [], []


# ── XLSX rendering ─────────────────────────────────────────────────────────────

def _render_xlsx_sheet(ws, variables: dict[str, Any]) -> None:
    """Two-pass render: expand loops first, then scalar substitution."""
    _expand_loops_xlsx(ws, variables)
    _substitute_scalars_xlsx(ws, variables)


def _expand_loops_xlsx(ws, variables: dict[str, Any]) -> None:
    """Find {{#var}} / {{/var}} marker rows and expand them in-place."""
    # Collect loop sections (scan all rows, first column)
    sections: list[tuple[int, int, str]] = []  # (start_row, end_row, var_name)

    row_count = ws.max_row
    i = 1
    while i <= row_count:
        val = _cell_str(ws.cell(row=i, column=1))
        ms = LOOP_START_RE.match(val)
        if ms:
            var_name = ms.group(1)
            for j in range(i + 1, row_count + 2):
                ev = _cell_str(ws.cell(row=j, column=1))
                me = LOOP_END_RE.match(ev)
                if me and me.group(1) == var_name:
                    sections.append((i, j, var_name))
                    i = j + 1
                    break
            else:
                i += 1
        else:
            i += 1

    # Process bottom-to-top so row indices stay valid
    for start_row, end_row, var_name in reversed(sections):
        items = variables.get(var_name)
        if not isinstance(items, list):
            items = []

        # Template rows are between start+1 and end-1 inclusive
        template_rows: list[list[Any]] = []
        max_col = ws.max_column
        for r in range(start_row + 1, end_row):
            row_vals = [ws.cell(row=r, column=c).value for c in range(1, max_col + 1)]
            template_rows.append(row_vals)

        num_tpl   = len(template_rows)
        num_items = len(items)
        num_insert = num_tpl * num_items

        insert_at = start_row + 1  # insert right after start marker

        # Insert blank rows for the expanded data
        if num_insert > 0:
            ws.insert_rows(insert_at, num_insert)
            for item_idx, item in enumerate(items):
                enriched = {"stt": str(item_idx + 1), **item}
                for tpl_offset, tpl_row in enumerate(template_rows):
                    actual_row = insert_at + item_idx * num_tpl + tpl_offset
                    for col_idx, val in enumerate(tpl_row, start=1):
                        new_val = _sub(str(val) if val is not None else "", {}, enriched)
                        ws.cell(row=actual_row, column=col_idx).value = new_val or None

        # After inserting num_insert rows at start_row+1:
        #   - start marker stays at start_row (insert was after it)
        #   - original template rows shifted to start_row + num_insert + 1
        #   - end marker shifted to end_row + num_insert
        actual_end       = end_row + num_insert
        actual_tpl_start = start_row + num_insert + 1

        ws.delete_rows(actual_end)                    # end marker
        ws.delete_rows(actual_tpl_start, num_tpl)     # original template rows
        ws.delete_rows(start_row)                     # start marker (not shifted)


def _substitute_scalars_xlsx(ws, variables: dict[str, Any]) -> None:
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and "{{" in cell.value:
                cell.value = _sub(cell.value, variables, None) or None


# ── DOCX rendering ─────────────────────────────────────────────────────────────

def _render_docx(doc, variables: dict[str, Any]) -> None:
    # 1. Process standalone paragraphs (scalar substitution + loop detection)
    _render_paragraphs(doc.paragraphs, variables)

    # 2. Process tables (loop expansion + scalar substitution)
    for table in doc.tables:
        _render_table(table, variables)

    # 3. Process headers/footers
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr:
                _render_paragraphs(hdr.paragraphs, variables)


def _render_paragraphs(paragraphs, variables: dict[str, Any]) -> None:
    for para in paragraphs:
        full = para.text
        if "{{" in full:
            # Rebuild runs with substituted text
            new_text = _sub(full, variables, None)
            if new_text != full:
                _set_paragraph_text(para, new_text)


def _render_table(table, variables: dict[str, Any]) -> None:
    """
    For tables: detect loop marker rows and expand them.
    A loop marker row has exactly one non-empty cell: {{#var}} or {{/var}}.
    """
    rows = table.rows
    sections: list[tuple[int, int, str]] = []
    n = len(rows)

    i = 0
    while i < n:
        val = _row_first_text(rows[i])
        ms = LOOP_START_RE.match(val)
        if ms:
            var_name = ms.group(1)
            for j in range(i + 1, n + 1):
                if j == n:
                    break
                ev = _row_first_text(rows[j])
                me = LOOP_END_RE.match(ev)
                if me and me.group(1) == var_name:
                    sections.append((i, j, var_name))
                    i = j + 1
                    break
            else:
                i += 1
        else:
            i += 1

    # Expand from bottom to top
    for start_idx, end_idx, var_name in reversed(sections):
        items = variables.get(var_name)
        if not isinstance(items, list):
            items = []

        template_rows = list(table.rows)[start_idx + 1:end_idx]

        # For each item, add new rows by copying the template
        # Insert after end_idx - 1 (just before end marker)
        insert_ref_row = table.rows[end_idx - 1]._tr

        for item in reversed(items):
            enriched = {"stt": str(items.index(item) + 1), **item}
            for tpl_row in reversed(template_rows):
                new_tr = copy.deepcopy(tpl_row._tr)
                insert_ref_row.addprevious(new_tr)
                # Substitute in the new row's cells
                from docx.table import _Row
                from lxml import etree
                for cell_el in new_tr.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"):
                    for p_el in cell_el.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                        texts = [r.text or "" for r in p_el.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")]
                        combined = "".join(texts)
                        if "{{" in combined:
                            new_txt = _sub(combined, {}, enriched)
                            t_els = p_el.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                            if t_els:
                                t_els[0].text = new_txt
                                for t in t_els[1:]:
                                    t.text = ""

        # Remove template rows and markers (in forward order, removing by _tr)
        tbl_el = table._tbl
        for r in list(table.rows)[start_idx:end_idx + 1]:
            tbl_el.remove(r._tr)

    # Scalar substitution for remaining rows
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full = para.text
                if "{{" in full:
                    new_text = _sub(full, variables, None)
                    if new_text != full:
                        _set_paragraph_text(para, new_text)


# ── Variable parser ────────────────────────────────────────────────────────────

def _parse_variables_xlsx(path: str) -> tuple[list[str], list[str]]:
    import openpyxl
    scalars: set[str] = set()
    lists:   set[str] = set()

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str) or "{{" not in cell.value:
                    continue
                for m in VAR_RE.finditer(cell.value):
                    key = m.group(1)
                    if key.startswith("item."):
                        continue
                    if LOOP_START_RE.match("{{" + key + "}}"):
                        lists.add(key[1:])   # strip leading #
                    elif not LOOP_END_RE.match("{{" + key + "}}"):
                        scalars.add(key)
    wb.close()

    # Also detect loop markers from full cell values
    wb2 = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for ws in wb2.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                v = str(cell.value or "")
                ms = LOOP_START_RE.match(v)
                if ms:
                    lists.add(ms.group(1))
    wb2.close()

    return sorted(scalars), sorted(lists)


def _parse_variables_docx(path: str) -> tuple[list[str], list[str]]:
    from docx import Document
    scalars: set[str] = set()
    lists:   set[str] = set()

    doc = Document(path)
    all_texts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_texts.append(cell.text)

    for text in all_texts:
        ms = LOOP_START_RE.match(text.strip())
        if ms:
            lists.add(ms.group(1))
            continue
        for m in VAR_RE.finditer(text):
            key = m.group(1)
            if not key.startswith("item."):
                scalars.add(key)

    return sorted(scalars), sorted(lists)


# ── PDF via LibreOffice ────────────────────────────────────────────────────────

def _convert_docx_to_pdf(docx_path: str, pdf_path: str) -> str:
    """Convert DOCX to PDF using LibreOffice headless. Returns actual output path."""
    out_dir = str(Path(pdf_path).parent)
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", out_dir, docx_path],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0:
            # LibreOffice names output as original filename with .pdf
            generated = Path(docx_path).with_suffix(".pdf")
            generated = Path(out_dir) / generated.name
            if generated.exists() and str(generated) != pdf_path:
                generated.rename(pdf_path)
            logger.info("PDF conversion OK: %s", pdf_path)
            return pdf_path
        else:
            logger.warning("LibreOffice conversion failed: %s", result.stderr)
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        logger.warning("LibreOffice not available (%s), returning DOCX", e)

    # Fallback: return DOCX path
    return docx_path


# ── Internal helpers ───────────────────────────────────────────────────────────

def _sub(text: str, scalars: dict[str, Any], item: dict | None) -> str:
    """Replace {{var}} and {{item.field}} in text."""
    if "{{" not in text:
        return text

    def replacer(m: re.Match) -> str:
        key = m.group(1)
        if item is not None and key.startswith("item."):
            field = key[5:]
            return str(item.get(field, ""))
        val = scalars.get(key)
        if val is None:
            return m.group(0)   # leave unknown vars intact
        if isinstance(val, list):
            return m.group(0)   # lists are handled by loop expansion
        return str(val)

    return VAR_RE.sub(replacer, text)


def _cell_str(cell) -> str:
    return str(cell.value or "").strip()


def _row_first_text(row) -> str:
    if not row.cells:
        return ""
    return row.cells[0].text.strip()


def _set_paragraph_text(para, new_text: str) -> None:
    """Replace all run text in a paragraph with new_text, keeping first run's formatting."""
    for i, run in enumerate(para.runs):
        run.text = new_text if i == 0 else ""
